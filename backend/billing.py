import sqlite3
import os
import sys
from datetime import datetime
import docx
from win32com.shell import shell, shellcon
from database.database import get_db_connection
from backend.inventory import get_product

def _get_base_path():
    """Gets the base path, for frozen applications."""
    if getattr(sys, 'frozen', False):
        # The application is running as a bundled executable
        return sys._MEIPASS
    else:
        # The application is running as a normal Python script
        # Go up one level from the 'backend' directory to the project root
        return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

def create_bill(items):
    """
    Creates a new bill, adds items to it, and updates product stock.
    - items: A list of dictionaries, where each dict contains:
             {'product_id': <id>, 'quantity': <qty>, 'price': <price>}
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    total_amount = 0

    try:
        for item in items:
            product = get_product(item['product_id'])
            if product['quantity'] < item['quantity']:
                raise ValueError(f"Not enough stock for {product['name']}.")
            total_amount += item['price'] * item['quantity']

        cursor.execute("INSERT INTO bills (total_amount) VALUES (?)", (total_amount,))
        bill_id = cursor.lastrowid

        for item in items:
            cursor.execute(
                "INSERT INTO bill_items (bill_id, product_id, quantity, price) VALUES (?, ?, ?, ?)",
                (bill_id, item['product_id'], item['quantity'], item['price'])
            )
            cursor.execute(
                "UPDATE products SET quantity = quantity - ? WHERE id = ?",
                (item['quantity'], item['product_id'])
            )

        conn.commit()
        return bill_id
    except (sqlite3.Error, ValueError) as e:
        conn.rollback()
        print(f"Error creating bill: {e}")
        return None
    finally:
        conn.close()

def get_bill_details(bill_id):
    """Retrieves details for a specific bill."""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM bills WHERE id = ?", (bill_id,))
    bill_info = cursor.fetchone()
    cursor.execute("""
        SELECT p.name, bi.quantity, bi.price
        FROM bill_items bi
        JOIN products p ON bi.product_id = p.id
        WHERE bi.bill_id = ?
    """, (bill_id,))
    bill_items = cursor.fetchall()
    conn.close()
    return bill_info, bill_items

def set_cell_text(cell, text):
    """
    Sets the text in a cell, clearing previous content and adding a new run
    to help preserve the cell's inherited paragraph style.
    """
    if not cell.paragraphs:
        cell.add_paragraph()
    
    p = cell.paragraphs[0]
    p.clear()
    p.add_run(text)

def fill_bill(template_path, output_path, items, date_str):
    """
    Fills a bill template with item data, placing the total in a fixed location
    and leaving the payment option column untouched.

    Args:
        template_path (str): Path to the .docx template.
        output_path (str): Path to save the filled .docx file.
        items (list): A list of dictionaries representing items.
        date_str (str): The date for the bill.
    """
    try:
        document = docx.Document(template_path)
        
        # --- Header/Summary Table (Table 1) ---
        header_table = document.tables[1]
        
        # Place date in its correct cell. The payment option cell (1,1) is intentionally left untouched.
        header_table.cell(1, 2).text = date_str

        # --- Items Table (Table 2) ---
        items_table = document.tables[2]

        # Set 'QTY' header text
        set_cell_text(items_table.cell(0, 0), "QTY")
        
        # --- Item Processing ---
        max_items = 15 
        if len(items) > max_items:
            print(f"Warning: Template supports {max_items} items. Ignoring the last {len(items) - max_items} items.")
            items_to_process = items[:max_items]
        else:
            items_to_process = items
        
        total = 0
        for i, item in enumerate(items_to_process):
            row = items_table.rows[i + 1]
            
            qty = item['qty']
            name = item['name']
            per_price = item['per-price']
            sub_total = qty * per_price
            total += sub_total
            
            row.cells[0].text = str(qty)
            row.cells[1].text = name
            row.cells[2].text = str(per_price)
            row.cells[3].text = str(sub_total)

        # --- Place Total in the fixed location (Row 16) ---
        total_row_index = 16
        if total_row_index < len(items_table.rows):
            total_row = items_table.rows[total_row_index]
            set_cell_text(total_row.cells[2], "TOTAL")
            set_cell_text(total_row.cells[3], str(total))
        else:
            print(f"Error: The template does not have the expected row {total_row_index} for the total.")

        document.save(output_path)
        print(f"Bill saved to {output_path}. Payment option column is now preserved.")
        return True # Indicate success
        
    except IndexError as e:
        print(f"Error: A table or cell index was out of range. Check template structure. Details: {e}")
        return False
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return False

def generate_bill_docx(bill_id, filename_suffix=""):
    bill_info, bill_items_raw = get_bill_details(bill_id)
    if not bill_info:
        print(f"No bill found with ID: {bill_id}")
        return None

    items_for_docx = [{'qty': item['quantity'], 'name': item['name'], 'per-price': item['price']} for item in bill_items_raw]
    date_str = datetime.strptime(bill_info['bill_date'], '%Y-%m-%d %H:%M:%S').strftime("%d-%b-%Y")

    documents_path = shell.SHGetFolderPath(0, shellcon.CSIDL_PERSONAL, None, 0)
    bills_dir = os.path.join(documents_path, 'bills')
    os.makedirs(bills_dir, exist_ok=True)
    
    output_path = os.path.join(bills_dir, f"bill_{bill_id}{filename_suffix}.docx")
    
    # Construct an absolute path to the template file using the helper function
    template_path = os.path.join(_get_base_path(), "Template.docx")

    try:
        if fill_bill(template_path, output_path, items_for_docx, date_str):
            return output_path
        else:
            return None
        
    except FileNotFoundError:
        print(f"Error: Template file not found at '{template_path}'. Please ensure it is in the root directory.")
        return None
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return None

def get_all_bills():
    """Retrieves all bills from the database."""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, bill_date, total_amount FROM bills ORDER BY bill_date DESC")
    bills = cursor.fetchall()
    conn.close()
    return bills